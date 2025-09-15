import os
import re
import json
import io
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import fitz
import docx
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import pytesseract

load_dotenv()

app = Flask(__name__)

CORS(app, resources={
    r"/*": {"origins": ["https://meu-agente-juridico.vercel.app", "http://localhost:5173"]}
})

# --- Funções de Extração de Texto ---
def extrair_texto_pdf(caminho_arquivo):
    try:
        doc = fitz.open(stream=caminho_arquivo.read(), filetype="pdf")
        return "".join(pagina.get_text() for pagina in doc)
    except Exception as e: return f"Erro PDF: {e}"

def extrair_texto_docx(caminho_arquivo):
    try:
        doc = docx.Document(caminho_arquivo)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e: return f"Erro DOCX: {e}"

def extrair_texto_txt(caminho_arquivo):
    try:
        return caminho_arquivo.read().decode('utf-8')
    except Exception as e: return f"Erro TXT: {e}"

def extrair_texto_imagem(caminho_arquivo):
    try:
        return pytesseract.image_to_string(Image.open(caminho_arquivo))
    except Exception as e: return f"Erro Imagem: {e}"

def limpar_texto_ia(texto):
    texto_limpo = texto.replace('**', '').replace('##', '').replace('`', '')
    if texto_limpo.strip().startswith("json"):
        texto_limpo = texto_limpo.strip()[4:]
    return texto_limpo

# --- Agentes de IA ---

# ETAPA 1: Gerar o rascunho com o Qwen
def etapa1_gerar_rascunho_ia(contexto, tipo_peticao):
    try:
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=os.getenv("OPENAI_API_KEY"),
        )
        
        prompt_sistema_gerador = """
        **IDENTIDADE E PERSONA:** Você é o Agente Peticionador do Escritório R. Feitosa Group. Sua persona é de um advogado sênior, erudito, crítico e inovador. Sua missão é redigir uma peça processual completa, com excelência técnica, argumentação robusta e fundamentação de nível acadêmico.

        **PROCESSO INTERNO DE GERAÇÃO E AUTO-REVISÃO CRÍTICA (ETAPAS OBRIGATÓRIAS):**
        Antes de escrever a palavra final, você DEVE seguir rigorosamente este processo interno:
        1.  **ANÁLISE ESTRATÉGICA:** Analise todo o contexto fornecido, confronte os fatos alegados com os documentos e identifique as teses jurídicas principais e subsidiárias.
        2.  **ELABORAÇÃO DO RASCUNHO:** Com base na análise, estruture e redija um rascunho completo da peça processual.
        3.  **AUTO-REVISÃO CRÍTICA (CHECKLIST DO REVISOR SÊNIOR):** Revise seu próprio rascunho de forma imparcial e metódica, verificando CADA um dos seguintes pontos:
            * **Clareza e Técnica:** O texto está objetivo, técnico e preciso? A linguagem está formal e sem jargões?
            * **Estrutura:** A peça contém todos os elementos obrigatórios (endereçamento, qualificação, fatos, direito, pedidos)? A estrutura está lógica e coerente?
            * **Técnica Jurídica:** As teses estão bem desenvolvidas? Os pedidos são claros, certos e determinados? Há pedidos subsidiários para aumentar a chance de êxito?
            * **Conformidade Legal:** As leis e artigos citados são pertinentes e foram transcritos e explicados corretamente?
            * **Jurisprudência:** A jurisprudência buscada é REAL, RECENTE (últimos 2 anos) e RELEVANTE para as teses? A ementa completa foi incluída?
            * **Fatos vs. Documentos:** Cada fato narrado está estritamente amparado pelo contexto e documentos fornecidos? **NÃO HÁ NENHUMA ALUCINAÇÃO?**
            * **Lacunas:** As informações ausentes foram devidamente apontadas no texto (ex: "[INFORMAÇÃO AUSENTE]")?
        4.  **VERSÃO FINAL:** Produza a versão final e polida da peça processual, incorporando todas as correções e melhorias da sua auto-revisão. **Sua saída deve conter APENAS esta versão final.**

        **REGRAS INVIOLÁVEIS (HARD GUARDS):**
        - **ZERO ALUCINAÇÃO:** Baseie-se **exclusivamente** nos fatos e documentos do contexto. JAMAIS invente fatos, leis, súmulas ou jurisprudência.
        - **PROLIXIDADE E PROFUNDIDADE:** Desenvolva cada tese com profundidade acadêmica.
        - **FUNDAMENTAÇÃO COMPLETA:** Transcreva leis, inclua ementas completas de jurisprudência e cite doutrina quando aplicável.
        - **FORMATO DE SAÍDA:** Gere APENAS o texto puro da petição, sem markdown, usando títulos em maiúsculas para as seções.
        """

        prompt_usuario = f"""
        Com base na sua identidade e regras, analise o contexto abaixo e redija uma **{tipo_peticao.replace('_', ' ').upper()}** completa, detalhada e prolixa. Desenvolva os argumentos e busque a fundamentação necessária (lei, jurisprudência, doutrina).

        **Contexto Extraído dos Documentos:**
        ---
        {contexto}
        ---
        """
        
        response = client.chat.completions.create(
            model="google/gemini-2.5-pro",
            messages=[
                {"role": "system", "content": prompt_sistema_gerador},
                {"role": "user", "content": prompt_usuario}
            ]
        )
        texto_bruto = response.choices[0].message.content
        return limpar_texto_ia(texto_bruto)
    except Exception as e:
        print(f"Erro na ETAPA 1 (Geração com Qwen): {e}")
        return f"Erro ao gerar o rascunho da IA. Detalhe: {e}"

# ETAPA 2: Revisar o rascunho com o GPT-4
def etapa2_revisar_relatorio_ia(texto_rascunho):
    try:
        client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=os.getenv("OPENAI_API_KEY"),
        )
        
        prompt_sistema_revisor = """
        Você atuará como um advogado peticionador sênior, um revisor jurídico imparcial, metódico e altamente criterioso. Sua função é revisar a peça processual fornecida.
        
        REGRAS DE REVISÃO:
        1.  **CORREÇÃO TÉCNICA:** Corrija ortografia, gramática e verifique a estrutura (elementos obrigatórios, coerência).
        2.  **ANÁLISE JURÍDICA:** Analise a robustez das teses, a adequação dos pedidos e verifique a conformidade com a legislação citada.
        3.  **VERIFICAÇÃO DE FATOS E FONTES:** Sua principal tarefa é garantir ZERO ALUCINAÇÃO. A peça deve se basear **exclusivamente** nos fatos do contexto original (que você não vê, mas o texto deve implicar). A jurisprudência citada deve ser real e relevante. Se algo parecer inventado, corrija ou remova.
        4.  **CLAREZA E OBJETIVIDADE:** Aprimore a redação para ser mais clara, objetiva e técnica, usando frases curtas e precisão terminológica, sem jargão popular.
        5.  **SAÍDA FINAL:** Sua única saída deve ser a versão final e reescrita da peça processual, incorporando todas as suas correções e melhorias. Não inclua comentários, pareceres ou checklists. Apenas o texto da petição corrigida.
        """

        prompt_usuario_revisor = f"""
        Revise a seguinte peça processual de acordo com suas regras. Entregue apenas a versão final e aprimorada do texto.

        **Peça para Revisão:**
        ---
        {texto_rascunho}
        ---
        """

        response = client.chat.completions.create(
            model="openai/gpt-5-nano",
            messages=[
                {"role": "system", "content": prompt_sistema_revisor},
                {"role": "user", "content": prompt_usuario_revisor}
            ]
        )
        texto_revisado = response.choices[0].message.content
        return limpar_texto_ia(texto_revisado)
    except Exception as e:
        print(f"Erro na ETAPA 2 (Revisão com GPT-4): {e}")
        return f"Erro na revisão da IA. Detalhe: {e}"


# --- Rotas da API ---
@app.route('/')
def health_check():
    return jsonify({"status": "ok", "message": "Servidor do Agente de Petição está no ar!"})

@app.route('/upload', methods=['POST'])
def upload_e_gerar_relatorio():
    if 'files' not in request.files:
        return jsonify({"error": "Nenhum arquivo enviado"}), 400

    files = request.files.getlist('files')
    tipo_peticao = request.form.get('tipo_peticao', 'Petição')
    contexto_completo, erros = "", []

    for file in files:
        if not file or not file.filename:
            continue
        filename = file.filename
        extensao = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        file_stream = file.stream
        texto_extraido = ""
        if extensao == 'pdf':
            texto_extraido = extrair_texto_pdf(file_stream)
        elif extensao == 'docx':
            texto_extraido = extrair_texto_docx(file_stream)
        elif extensao == 'txt':
            texto_extraido = extrair_texto_txt(file_stream)
        elif extensao in ['png', 'jpg', 'jpeg']:
            texto_extraido = extrair_texto_imagem(file_stream)
        else:
            erros.append(f"Formato não suportado: {filename}")
            continue
        contexto_completo += f"\n\n--- Conteúdo de: {filename} ---\n{texto_extraido}"

    if not contexto_completo and not erros:
        return jsonify({"error": "Não foi possível extrair conteúdo dos arquivos.", "details": erros}), 400

    # --- CADEIA DE CHAMADAS DE IA ---
    print("Iniciando Etapa 1: Geração do Rascunho com Qwen...")
    rascunho = etapa1_gerar_rascunho_ia(contexto_completo, tipo_peticao)
    
    if "Erro ao gerar" in rascunho:
        return jsonify({"error": rascunho, "stage": "generation"}), 500

    print("Iniciando Etapa 2: Revisão do Rascunho com GPT-4...")
    versao_final = etapa2_revisar_relatorio_ia(rascunho)

    if "Erro na revisão" in versao_final:
        return jsonify({"error": versao_final, "stage": "review"}), 500

    print("Processo concluído com sucesso!")
    return jsonify({"message": "Relatório gerado e revisado com sucesso!", "report": versao_final, "errors": erros})


@app.route('/download-report', methods=['POST'])
def download_report():
    # ... (seu código de download continua o mesmo, sem alterações) ...
    data = request.get_json()
    report_text = data.get('report')

    if not report_text:
        return jsonify({"error": "Nenhum texto de relatório fornecido"}), 400

    document = Document()
    
    style = document.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(14)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5

    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.05)
        section.right_margin = Cm(2.54)

    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            document.add_paragraph()
            continue
        
        p = document.add_paragraph()
        p.text = line
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(3.05)
        
        if (line.isupper() and len(line.split()) < 7) or re.match(r'^\d+\.\d+\.', line) or re.match(r'^[IVX]+\s?–', line):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(12)
        
        elif line.startswith("EXCELENTÍSSIMO"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            p.paragraph_format.first_line_indent = None

        elif line.startswith("Termos em que,") or line.startswith("ADVOGADO"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None

    document.add_paragraph()
    fecho1 = document.add_paragraph(f'Sobral/CE, na data da assinatura digital.')
    fecho1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecho1.paragraph_format.first_line_indent = None
    fecho2 = document.add_paragraph('Karlos Roneely Rocha Feitosa')
    fecho2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecho2.runs[0].bold = True
    fecho2.paragraph_format.first_line_indent = None
    fecho3 = document.add_paragraph('OAB-CE 23.104')
    fecho3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecho3.paragraph_format.first_line_indent = None
    
    file_stream_in_memory = io.BytesIO()
    document.save(file_stream_in_memory)
    file_stream_in_memory.seek(0)

    return send_file(
        file_stream_in_memory,
        as_attachment=True,
        download_name='peticao_gerada.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    app.run(debug=True, port=5001)